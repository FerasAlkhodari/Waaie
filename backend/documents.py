"""Structure-aware, page-aware text extraction for uploaded study documents.

Supports the three formats students attach — PDF, Word (.docx) and Excel
(.xlsx). PDFs are parsed *positionally* with ``pdfplumber`` (MIT, built on
pdfminer.six): every physical page is reconstructed in true reading order
(multi-column band ordering + per-line RTL/LTR ordering) and wrapped in
machine-parseable page-boundary tags so the DeepSeek mentor can anchor an
answer to a real page number instead of hallucinating across pages. Word and
Excel are rendered as light Markdown (headings, bullet lists, simple tables).

A 150,000-character context budget plus a *tag-aware* compressor let students
upload complete chapters while still guarding against runaway token spend: when
a document is over budget, page *bodies* are truncated proportionally but no
page-boundary tag is ever dropped, so every page stays addressable.

The scope guardrails in ``model.SYSTEM_INSTRUCTION`` still decide whether a
document is in scope — no subject filtering happens here.

Statelessness: extraction is a pure, deterministic function of the input bytes.
There is no database, cache, temp file, or module-level mutable state — the same
bytes in always produce an identical string out.
"""

import io
import logging
import re
import statistics
import unicodedata

import pdfplumber
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook

logger = logging.getLogger(__name__)

# Context budget (characters). On the prepaid DeepSeek production tier the
# model handles large contexts efficiently, so this is generous enough for a
# complete, detailed study document. Only content beyond this is compressed
# (structurally for Word/Excel, tag-aware per page for PDF) rather than
# hard-truncated, and the final string is guaranteed to be <= this many chars.
MAX_CONTEXT_CHARS = 150_000

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".xlsx")

# --------------------------------------------------------------------------- #
# Page-boundary tag schema — single source of truth.
# --------------------------------------------------------------------------- #
#
# Every physical page ``n`` is wrapped ``PAGE_START_TMPL(n)`` ... body ...
# ``PAGE_END_TMPL(n)``. Tests import these constants instead of hard-coding the
# strings, and every tag is matchable with the single deterministic regex
# ``PAGE_TAG_RE`` below — so downstream tooling (and any model) can recover the
# exact page boundaries with one pass.
PAGE_START_TMPL = "--- [PAGE {n} START] ---"
PAGE_END_TMPL = "--- [PAGE {n} END] ---"
PAGE_TRUNC_TMPL = "--- [PAGE {n} TRUNCATED: kept {kept} of {total} chars] ---"
PAGE_EMPTY_TMPL = "--- [PAGE {n}: NO EXTRACTABLE TEXT] ---"

# Matches a START or END boundary tag and captures (page_number, kind). Anchored
# and single-line so it only fires on a tag that occupies a whole line.
PAGE_TAG_RE = re.compile(r"^--- \[PAGE (\d+) (START|END)\] ---$")

# Appended (in Arabic) when a Word/Excel document had to be compressed, so the
# student knows the answer is based on a reduced view. PDFs carry their own
# per-page truncation markers instead, so this notice is not appended to them.
_COMPRESSION_NOTICE = (
    "\n\n> ملاحظة: هذا المستند كبير وتم اختصاره تلقائيًا إلى العناوين والفقرات "
    "الأساسية لتوفير سياق فعّال. للحصول على إجابات أدق، ارفع ملفًا أصغر خاصًا "
    "بالفصل أو الدرس المطلوب."
)


class DocumentParseError(Exception):
    """Raised when an upload cannot be read as a supported text document."""


class DocumentTooLargeError(DocumentParseError):
    """Raised when a document has so many pages that the page-boundary tags
    *alone* exceed the context budget — there is no room left for any body, so
    pages cannot be made addressable. Subclasses ``DocumentParseError`` so the
    API surfaces it as a clean 400 like any other unreadable upload."""


def extract_text(filename: str, data: bytes) -> str:
    """Extract clean, structure-aware text from a PDF/.docx/.xlsx payload.

    PDFs come back page-aware: every physical page is wrapped in START/END
    boundary tags and reconstructed in reading order. The returned string is
    always <= ``MAX_CONTEXT_CHARS``.

    Raises ``DocumentParseError`` for unsupported types, unreadable files, or
    documents that contain no extractable text (e.g. a fully scanned PDF), and
    ``DocumentTooLargeError`` when page tags alone blow the budget.
    """
    name = (filename or "").lower()

    if name.endswith(".pdf"):
        # PDF owns its full pipeline (positional parse + tag-aware budget) and
        # returns the final, already-budgeted string.
        return _extract_pdf(data)
    elif name.endswith(".docx"):
        markdown = _extract_docx(data)
    elif name.endswith(".xlsx"):
        markdown = _extract_xlsx(data)
    else:
        raise DocumentParseError(
            "Unsupported file type. Please upload a .pdf, .docx or .xlsx file."
        )

    markdown = _normalize(markdown)
    if not markdown:
        raise DocumentParseError(
            "No readable text was found in the document. If it is a scanned "
            "PDF (images only), it cannot be parsed."
        )

    return _compress_to_budget(markdown, MAX_CONTEXT_CHARS)


# --------------------------------------------------------------------------- #
# PDF — positional, structure-aware, page-aware extraction
# --------------------------------------------------------------------------- #
#
# Library decision (see waaie_pdf_pipeline_prompt.md, Constraint 4): we use
# pdfplumber (MIT, on top of pdfminer.six). It exposes char/word-level bounding
# boxes via ``page.extract_words(...)`` which is exactly what the line-grouping,
# column-clustering and RTL-ordering algorithm below needs. We deliberately do
# NOT use PyMuPDF/fitz: it is AGPL-3.0, and for a hosted SaaS that students
# reach over the network AGPL triggers a source-disclosure obligation whose
# only exemption is Artifex's *paid* commercial license — both of which violate
# the project's licensing constraint. PyMuPDF would only be acceptable if the
# team explicitly accepted open-sourcing the backend under AGPL or buying the
# commercial license; absent that flag we default to pdfplumber.
#
# OCR is out of scope: image-only/scanned pages have no text layer and are
# represented with PAGE_EMPTY_TMPL rather than sent to any cloud OCR service. A
# future, fully-local option would be pytesseract + Tesseract, but it is not
# implemented here.

# Arabic script blocks (per the brief): base Arabic, Supplement, Extended-A, and
# the two Presentation-Forms blocks (which is what reshaped PDF text decodes to
# before NFKC normalization folds it back to base letters).
_ARABIC_SCRIPT_RE = re.compile(
    "[؀-ۿݐ-ݿࢠ-ࣿﭐ-﷿ﹰ-﻿]"
)
_LATIN_RE = re.compile(r"[A-Za-z]")


def _is_rtl(text: str) -> bool:
    """A run is RTL-dominant when its Arabic-script characters outnumber its
    strong-LTR (Latin) characters. Neutral characters (digits, punctuation,
    whitespace) do not vote, so a line of pure numbers stays LTR."""
    return len(_ARABIC_SCRIPT_RE.findall(text)) > len(_LATIN_RE.findall(text))


def _median(values, default: float) -> float:
    return statistics.median(values) if values else default


def _detect_column_bands(words, page_width: float):
    """Cluster ``words`` into vertical column bands separated by whitespace
    gutters, returned left-to-right.

    A gutter is an x-range that *no* word crosses anywhere on the page and that
    is wider than a threshold derived from the median character width (and a
    floor relative to page width). Because every word's x-interval is merged,
    an element that physically spans two columns bridges the gutter and collapses
    them into one band — the conservative choice (we would rather under-split
    than interleave a real two-column layout). Within-line gaps never create a
    false column because other lines cover that x-range.
    """
    char_widths = [
        (w["x1"] - w["x0"]) / max(len(w["text"]), 1)
        for w in words
        if w["x1"] > w["x0"]
    ]
    cw = _median(char_widths, 5.0)
    gutter = max(page_width * 0.045, cw * 3.0)

    intervals = sorted((w["x0"], w["x1"]) for w in words)
    segments = []  # merged [x0, x1] covered runs, ascending
    for x0, x1 in intervals:
        if segments and x0 <= segments[-1][1] + gutter:
            if x1 > segments[-1][1]:
                segments[-1][1] = x1
        else:
            segments.append([x0, x1])

    if len(segments) <= 1:
        return [words]

    bands = [[] for _ in segments]
    for w in words:
        cx = (w["x0"] + w["x1"]) / 2
        placed = False
        for i, (s0, s1) in enumerate(segments):
            if s0 - 1e-6 <= cx <= s1 + 1e-6:
                bands[i].append(w)
                placed = True
                break
        if not placed:
            # Falls inside a gutter (rare): attach to the nearest band by center.
            nearest = min(
                range(len(segments)),
                key=lambda i: abs(cx - (segments[i][0] + segments[i][1]) / 2),
            )
            bands[nearest].append(w)

    return [b for b in bands if b]


def _group_lines(words):
    """Cluster a band's words into text lines by vertical position.

    Words whose ``top`` falls within a tolerance of the current line's anchor
    (derived from the median glyph height) join that line; otherwise a new line
    starts. Returns lines top-to-bottom, each a list of word dicts.
    """
    heights = [w["bottom"] - w["top"] for w in words if w["bottom"] > w["top"]]
    median_height = _median(heights, 10.0)
    y_tol = median_height * 0.6

    ordered = sorted(words, key=lambda w: (w["top"], w["x0"]))
    lines = []
    current = [ordered[0]]
    anchor = ordered[0]["top"]
    for w in ordered[1:]:
        if abs(w["top"] - anchor) <= y_tol:
            current.append(w)
        else:
            lines.append(current)
            current = [w]
            anchor = w["top"]
    lines.append(current)
    return lines, median_height


def _order_line(line_words) -> str:
    """Order one line's words into reading order and join with single spaces.

    Direction is detected per line from its own characters: an RTL-dominant line
    is read right→left (descending x), an LTR line left→right. Word text itself
    is never reversed, so Latin/digit runs keep their internal order (a year
    like ``2024`` never becomes ``4202``); this realizes Unicode-bidi-correct
    ordering at the line level.
    """
    text = "".join(w["text"] for w in line_words)
    if _is_rtl(text):
        ordered = sorted(line_words, key=lambda w: -w["x0"])
    else:
        ordered = sorted(line_words, key=lambda w: w["x0"])
    return " ".join(w["text"] for w in ordered)


def _reconstruct_page(page) -> dict:
    """Reconstruct one physical page into reading order.

    Returns a structural record: ``columns`` (band count), ``reading_order``
    (how bands were traversed), ``lines`` (reconstructed text lines),
    ``direction`` (page-dominant script), ``char_count`` and the normalized
    ``body`` string ('' when the page has no extractable text layer).
    """
    words = page.extract_words(use_text_flow=False, keep_blank_chars=False)
    if not words:
        return {
            "columns": 0,
            "reading_order": "none",
            "lines": [],
            "direction": "ltr",
            "char_count": 0,
            "body": "",
        }

    page_text = "".join(w["text"] for w in words)
    page_rtl = _is_rtl(page_text)

    bands = _detect_column_bands(words, page.width or 1.0)
    # bands come back left-to-right; an RTL document reads the rightmost first.
    if page_rtl:
        bands = list(reversed(bands))
        reading_order = "right-to-left" if len(bands) > 1 else "rtl-single-column"
    else:
        reading_order = "left-to-right" if len(bands) > 1 else "ltr-single-column"

    lines = []
    for band in bands:
        line_groups, median_height = _group_lines(band)
        para_gap = median_height * 1.6
        prev_bottom = None
        for group in line_groups:
            top = min(w["top"] for w in group)
            bottom = max(w["bottom"] for w in group)
            # A large vertical gap is a structural break (paragraph / question
            # boundary): keep it so text never merges across it.
            if prev_bottom is not None and (top - prev_bottom) > para_gap:
                lines.append("")
            lines.append(_order_line(group))
            prev_bottom = bottom

    body = _normalize("\n".join(lines))
    return {
        "columns": len(bands),
        "reading_order": reading_order,
        # Normalized lines (base Arabic letters, math as ASCII) so the trace
        # shows the same canonical text the model receives, in reading order.
        "lines": body.split("\n") if body else [],
        "direction": "rtl" if page_rtl else "ltr",
        "char_count": len(body),
        "body": body,
    }


def _parse_pdf_pages(data: bytes):
    """Open a PDF from in-memory bytes and reconstruct every page.

    Returns the list of per-page structural records (see ``_reconstruct_page``)
    in physical page order. Raises ``DocumentParseError`` if the byte stream is
    not a readable PDF (corrupt/truncated upload).
    """
    try:
        pages = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                pages.append(_reconstruct_page(page))
        return pages
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"Could not read the PDF file: {exc}")


def _fit_body(page_number: int, body: str, alloc: int):
    """Shrink one page body to fit ``alloc`` characters.

    Truncation keeps the *head* (so leading structural content — headers and
    question numbers — survives) and appends ``PAGE_TRUNC_TMPL`` recording how
    much was kept. The returned body is always <= ``alloc`` characters, which is
    what guarantees the document-level budget postcondition.
    """
    total = len(body)
    if total <= alloc:
        return body, False, total

    marker_max = len(
        PAGE_TRUNC_TMPL.format(n=page_number, kept=total, total=total)
    )
    head_room = max(alloc - marker_max - 1, 0)  # -1 for the newline before marker
    head = body[:head_room]
    marker = PAGE_TRUNC_TMPL.format(n=page_number, kept=len(head), total=total)
    out = f"{head}\n{marker}" if head else marker
    return out, True, len(head)


def _assemble_pdf_with_budget(bodies, budget: int):
    """Wrap each page body in START/END tags and assemble a single string that
    is guaranteed <= ``budget`` characters *including every tag*.

    Tag overhead (all START + END tags + structural newlines + inter-page
    joiners) is fixed and never reduced — no page-boundary tag is ever dropped,
    so every page number stays addressable. The remaining body budget is shared
    across pages proportionally to original page length, with a per-page floor
    big enough to always fit a truncation marker. Empty pages emit
    ``PAGE_EMPTY_TMPL``. If the tag overhead alone leaves no room for the page
    bodies/markers, ``DocumentTooLargeError`` is raised rather than silently
    dropping pages.

    Returns ``(result_string, page_meta)`` where ``page_meta[i]`` carries
    ``tag_start_offset``, ``tag_end_offset``, ``truncated``, ``kept`` and
    ``total`` for the structural parse trace.
    """
    n = len(bodies)
    starts = [PAGE_START_TMPL.format(n=i + 1) for i in range(n)]
    ends = [PAGE_END_TMPL.format(n=i + 1) for i in range(n)]

    # Fixed, never-reduced overhead: per page = start + "\n" + "\n" + end, plus a
    # "\n\n" joiner between consecutive pages.
    fixed_overhead = sum(len(starts[i]) + len(ends[i]) + 2 for i in range(n))
    fixed_overhead += 2 * max(n - 1, 0)
    body_budget = budget - fixed_overhead

    final_bodies = [None] * n
    nonempty = []
    empty_cost = 0
    for i, body in enumerate(bodies):
        if body:
            nonempty.append(i)
        else:
            marker = PAGE_EMPTY_TMPL.format(n=i + 1)
            final_bodies[i] = marker
            empty_cost += len(marker)

    # Pathological: tags (+ mandatory empty markers) alone exceed the budget.
    if body_budget < empty_cost:
        raise DocumentTooLargeError(
            f"This document has too many pages ({n}); its page-boundary tags "
            f"alone exceed the {budget:,}-character context budget. Please "
            f"upload a smaller file (a single chapter or lesson)."
        )

    avail = body_budget - empty_cost
    total_len = sum(len(bodies[i]) for i in nonempty)

    page_meta = [None] * n

    if total_len <= avail:
        # Everything fits — no body truncation needed.
        for i in nonempty:
            final_bodies[i] = bodies[i]
            page_meta[i] = {"truncated": False, "kept": len(bodies[i]),
                            "total": len(bodies[i])}
    else:
        # Reserve a per-page floor so each page can always hold either its whole
        # (tiny) body or a full truncation marker; distribute the rest in
        # proportion to original length.
        floor = {}
        for i in nonempty:
            length = len(bodies[i])
            marker_max = len(PAGE_TRUNC_TMPL.format(n=i + 1, kept=length, total=length))
            floor[i] = min(length, marker_max + 1)
        reserved = sum(floor.values())
        if reserved > avail:
            raise DocumentTooLargeError(
                f"This document is too large: even a minimal, truncated view of "
                f"its {n} pages exceeds the {budget:,}-character context budget. "
                f"Please upload a smaller file (a single chapter or lesson)."
            )

        extra = avail - reserved
        rem_len = sum(len(bodies[i]) - floor[i] for i in nonempty)
        for i in nonempty:
            length = len(bodies[i])
            share = int(extra * (length - floor[i]) / rem_len) if rem_len > 0 else 0
            alloc = floor[i] + share
            fitted, truncated, kept = _fit_body(i + 1, bodies[i], alloc)
            final_bodies[i] = fitted
            page_meta[i] = {"truncated": truncated, "kept": kept, "total": length}

    for i in range(n):
        if page_meta[i] is None:  # empty page
            page_meta[i] = {"truncated": False, "kept": 0, "total": 0}

    # Assemble, tracking each tag's offset for the structural parse trace.
    blocks = []
    offset = 0
    for i in range(n):
        block = f"{starts[i]}\n{final_bodies[i]}\n{ends[i]}"
        if i > 0:
            offset += 2  # the "\n\n" joiner that precedes this block
        page_meta[i]["tag_start_offset"] = offset
        page_meta[i]["tag_end_offset"] = offset + len(block) - len(ends[i])
        offset += len(block)
        blocks.append(block)

    result = "\n\n".join(blocks)
    # Postcondition: the full string, tags included, never exceeds the budget.
    assert len(result) <= budget, (
        f"budget postcondition violated: {len(result)} > {budget}"
    )
    return result, page_meta


def _extract_pdf(data: bytes) -> str:
    """Positional PDF extraction → tag-wrapped, budget-bounded string.

    Each physical page is reconstructed in reading order and wrapped in strict,
    sequential 1..N START/END tags that map directly to the page numbers a PDF
    viewer shows. A page with no text layer keeps its tags but carries
    ``PAGE_EMPTY_TMPL``, so numbering never drifts. If NO page yields any text
    (a fully scanned PDF), ``DocumentParseError`` is raised so the caller returns
    the existing "no readable text" message.
    """
    pages = _parse_pdf_pages(data)
    if not pages or all(not p["body"] for p in pages):
        raise DocumentParseError(
            "No readable text was found in the document. If it is a scanned "
            "PDF (images only), it cannot be parsed."
        )

    bodies = [p["body"] for p in pages]
    result, _ = _assemble_pdf_with_budget(bodies, MAX_CONTEXT_CHARS)
    return result


def structural_parse_trace(data: bytes) -> dict:
    """Build a deterministic structural parse tree for a PDF — the verification
    "AST" analogue. Returns ``document -> pages[{page, columns, reading_order,
    direction, lines, char_count, tag_start_offset, tag_end_offset, truncated,
    kept, total}]`` plus document-level ``page_count``,
    ``total_chars_pre_compression``, ``total_chars_post_compression`` and
    ``within_budget``. Pure and offline; used by the verification trace script.
    """
    pages = _parse_pdf_pages(data)
    bodies = [p["body"] for p in pages]
    pre = sum(len(b) for b in bodies)

    result, page_meta = _assemble_pdf_with_budget(bodies, MAX_CONTEXT_CHARS)

    page_nodes = []
    for i, parsed in enumerate(pages):
        meta = page_meta[i]
        page_nodes.append(
            {
                "page": i + 1,
                "columns": parsed["columns"],
                "reading_order": parsed["reading_order"],
                "direction": parsed["direction"],
                "lines": parsed["lines"],
                "char_count": parsed["char_count"],
                "tag_start_offset": meta["tag_start_offset"],
                "tag_end_offset": meta["tag_end_offset"],
                "truncated": meta["truncated"],
                "kept": meta["kept"],
                "total": meta["total"],
            }
        )

    return {
        "page_count": len(pages),
        "total_chars_pre_compression": pre,
        "total_chars_post_compression": len(result),
        "within_budget": len(result) <= MAX_CONTEXT_CHARS,
        "pages": page_nodes,
    }


# --------------------------------------------------------------------------- #
# Word (.docx) extraction
# --------------------------------------------------------------------------- #


def _extract_docx(data: bytes) -> str:
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise DocumentParseError(f"Could not read the Word document: {exc}")

    blocks = []
    for item in _iter_block_items(document):
        if isinstance(item, Paragraph):
            rendered = _paragraph_to_md(item)
        else:  # Table
            rendered = _table_to_md(item)
        if rendered:
            blocks.append(rendered)

    return "\n\n".join(blocks)


def _iter_block_items(document):
    """Yield paragraphs and tables in document order (python-docx loses order
    when iterating ``.paragraphs`` and ``.tables`` separately)."""
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _paragraph_to_md(paragraph: Paragraph) -> str:
    text = paragraph.text.strip()
    if not text:
        return ""

    style = (paragraph.style.name if paragraph.style else "") or ""
    style_low = style.lower()

    if style_low.startswith("title"):
        return f"# {text}"

    if style_low.startswith("heading"):
        digits = "".join(ch for ch in style_low if ch.isdigit())
        level = min(int(digits), 6) if digits else 2
        return f"{'#' * level} {text}"

    if "list" in style_low or paragraph._p.find(qn("w:numPr")) is not None:
        return f"- {text}"

    return text


def _table_to_md(table: Table) -> str:
    rows = []
    for row in table.rows:
        cells = [
            cell.text.strip().replace("\n", " ").replace("|", "\\|")
            for cell in row.cells
        ]
        if any(cells):
            rows.append(cells)

    if not rows:
        return ""

    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]

    lines = ["| " + " | ".join(rows[0]) + " |"]
    lines.append("| " + " | ".join(["---"] * width) + " |")
    for r in rows[1:]:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


# --------------------------------------------------------------------------- #
# Excel (.xlsx) extraction
# --------------------------------------------------------------------------- #


def _extract_xlsx(data: bytes) -> str:
    """Render every worksheet as a Markdown table. ``data_only=True`` returns
    computed cell values (not formulas) so DeepSeek can reason over the numbers
    directly; ``read_only=True`` keeps memory flat on large workbooks."""
    try:
        workbook = load_workbook(
            io.BytesIO(data), read_only=True, data_only=True
        )
    except Exception as exc:
        raise DocumentParseError(f"Could not read the Excel file: {exc}")

    blocks = []
    try:
        for sheet in workbook.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                cells = [_cell_to_text(value) for value in row]
                if any(cell for cell in cells):
                    rows.append(cells)

            if not rows:
                continue

            width = max(len(r) for r in rows)
            rows = [r + [""] * (width - len(r)) for r in rows]

            blocks.append(f"## Sheet: {sheet.title}")
            table = ["| " + " | ".join(rows[0]) + " |"]
            table.append("| " + " | ".join(["---"] * width) + " |")
            for r in rows[1:]:
                table.append("| " + " | ".join(r) + " |")
            blocks.append("\n".join(table))
    finally:
        workbook.close()

    return "\n\n".join(blocks)


def _cell_to_text(value) -> str:
    if value is None:
        return ""
    return str(value).strip().replace("\n", " ").replace("|", "\\|")


# --------------------------------------------------------------------------- #
# Whitespace / Unicode normalization & Word/Excel budget compression
# --------------------------------------------------------------------------- #


# Math glyphs that PDF/Word exports emit as Unicode but that get mangled or
# read poorly downstream. We translate them to plain-ASCII math equivalents so
# formulas survive extraction intact and stay legible to the model — e.g.
# "E = I² R t" -> "E = I^2 R t", "H₂O" -> "H_2O", "a × b − c" -> "a * b - c".
# Plain operators (= ^ + * / -) are already ASCII and pass through untouched.
# This MUST run before NFKC normalization, which would otherwise fold a
# superscript "²" into a plain "2" and lose the exponent semantics.
_MATH_TRANSLATIONS = {
    # Superscript digits -> ^n
    0x2070: "^0", 0x00B9: "^1", 0x00B2: "^2", 0x00B3: "^3", 0x2074: "^4",
    0x2075: "^5", 0x2076: "^6", 0x2077: "^7", 0x2078: "^8", 0x2079: "^9",
    0x207A: "^+", 0x207B: "^-", 0x207F: "^n",
    # Subscript digits -> _n (also helps chemistry: H2O, CO2)
    0x2080: "_0", 0x2081: "_1", 0x2082: "_2", 0x2083: "_3", 0x2084: "_4",
    0x2085: "_5", 0x2086: "_6", 0x2087: "_7", 0x2088: "_8", 0x2089: "_9",
    0x208A: "_+", 0x208B: "_-",
    # Operators that have unambiguous ASCII forms
    0x00D7: "*",   # × multiplication sign
    0x22C5: "*",   # ⋅ dot operator
    0x00B7: "*",   # · middle dot
    0x2217: "*",   # ∗ asterisk operator
    0x00F7: "/",   # ÷ division sign
    0x2215: "/",   # ∕ division slash
    0x2212: "-",   # − minus sign
    0x2260: "!=",  # ≠
    0x2264: "<=",  # ≤
    0x2265: ">=",  # ≥
}


# --------------------------------------------------------------------------- #
# Visual-order Arabic correction (string-level fallback)
# --------------------------------------------------------------------------- #
#
# The positional PDF parser above reconstructs Arabic word order from on-page
# coordinates, so it no longer relies on this string-level reversal heuristic.
# The helper is retained because (a) it is independently unit-tested and (b) it
# remains a safe, conservative repair for any path that hands us a line already
# flattened into visual (reversed) order. It is keyed on Arabic orthography so
# it distinguishes visual from logical order WITHOUT a dictionary and stays
# idempotent (a corrected line scores as logical and is never flipped again):
#   * teh marbuta (ة) and alef maksura (ى) can only END a logical word; seeing
#     one at the START of a token is a strong "this token is reversed" signal.
#   * the definite article "ال" begins many logical words; reversed it becomes a
#     trailing "لا", another reversal signal.
# A line is only flipped when reversal signals strictly outnumber logical ones.

_ARABIC_RE = re.compile(r"[؀-ۿ]")

# Characters stripped from a token's ends before inspecting its first/last
# letter (punctuation, quotes, digits, markdown bullets, whitespace).
_TOKEN_EDGE_STRIP = " \t\r\n.,:;!?()[]{}\"'«»،؛؟…-–—*#|/\\0123456789"

_TEH_MARBUTA = "ة"   # ة — final-only letter
_ALEF_MAKSURA = "ى"  # ى — final-only letter
_ALEF = "ا"          # ا
_LAM = "ل"           # ل
_DEFINITE_ARTICLE = _ALEF + _LAM   # "ال" (logical, word start)
_REVERSED_ARTICLE = _LAM + _ALEF   # "لا" (visual, word end)

# Latin/digit runs (incl. decimals/times like "3.14", "12:30") that must keep
# their internal left-to-right order when the surrounding RTL line is reversed —
# so "2024" never becomes "4202".
_LTR_RUN_RE = re.compile(r"[A-Za-z0-9]+(?:[.,:][A-Za-z0-9]+)*")


def _directional_scores(line: str) -> tuple[int, int]:
    """Return (logical_signals, reversed_signals) for the Arabic tokens in
    ``line`` based on the orthography heuristics described above."""
    logical = 0
    reversed_ = 0
    for token in line.split():
        if not _ARABIC_RE.search(token):
            continue
        core = token.strip(_TOKEN_EDGE_STRIP)
        if not core:
            continue
        if core.startswith(_DEFINITE_ARTICLE):
            logical += 1
        if core[-1] in (_TEH_MARBUTA, _ALEF_MAKSURA):
            logical += 1
        if core[0] in (_TEH_MARBUTA, _ALEF_MAKSURA):
            reversed_ += 1
        if len(core) > 2 and core.endswith(_REVERSED_ARTICLE):
            reversed_ += 1
    return logical, reversed_


def _reverse_visual_line(line: str) -> str:
    """Reverse a visually-ordered line back to logical order, restoring each
    Latin/digit run to its original left-to-right sequence."""
    reversed_chars = line[::-1]
    return _LTR_RUN_RE.sub(lambda m: m.group(0)[::-1], reversed_chars)


def _fix_visual_order_arabic(text: str) -> str:
    """Re-derive logical-order Arabic for any line a PDF emitted in visual order.

    Lines with no Arabic (\\u0600-\\u06FF) are returned unchanged. A line is only
    transformed when reversal signals strictly outnumber logical-order signals;
    otherwise it is left as-is. Idempotent: a corrected line scores as logical,
    so a second pass is a no-op."""
    if not text:
        return text

    fixed_lines = []
    changed = 0
    for line in text.split("\n"):
        if _ARABIC_RE.search(line):
            logical, reversed_ = _directional_scores(line)
            if reversed_ > logical and reversed_ >= 1:
                repaired = _reverse_visual_line(line)
                if repaired != line:
                    fixed_lines.append(repaired)
                    changed += 1
                    continue
        fixed_lines.append(line)

    if changed:
        logger.debug(
            "Corrected visual-order (reversed) Arabic on %d line(s).", changed
        )
    return "\n".join(fixed_lines)


def _normalize(text: str) -> str:
    """Collapse runs of spaces/tabs and blank lines, and canonicalize Unicode.

    The cheapest, safest token reduction: remove the dead whitespace PDF/Word
    exports are full of without dropping content. Math symbols are explicitly
    preserved — Unicode super/subscripts and operators map to readable ASCII
    (see ``_MATH_TRANSLATIONS``) rather than being dropped — and that mapping
    runs BEFORE ``NFKC`` so an exponent like "²" becomes "^2" instead of a flat
    "2". NFKC then folds Arabic presentation forms (what reshaped PDF text
    decodes to) back to their canonical base letters and splits the lam-alef
    ligature, giving the model clean, consistently-encoded Arabic.
    """
    text = text.translate(_MATH_TRANSLATIONS)
    text = unicodedata.normalize("NFKC", text)
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    collapsed = "\n".join(lines)
    collapsed = re.sub(r"\n{3,}", "\n\n", collapsed)
    return collapsed.strip()


def _compress_to_budget(text: str, budget: int) -> str:
    """If ``text`` fits the budget, return it untouched. Otherwise keep the
    document's skeleton — its leading intro, every heading, and the first
    paragraph under each heading — assembled in original order until the
    budget is spent, then append a notice nudging smaller uploads.

    Used for Word/Excel output (which has no physical pages). PDFs use the
    tag-aware ``_assemble_pdf_with_budget`` instead, which preserves page tags.
    """
    if len(text) <= budget:
        return text

    blocks = [b for b in text.split("\n\n") if b.strip()]
    is_heading = [b.lstrip().startswith("#") for b in blocks]
    keep = [False] * len(blocks)

    # Leading introduction (first few blocks set the document's context).
    for i in range(min(3, len(blocks))):
        keep[i] = True

    # Every heading plus the first block beneath it (its core paragraph/table).
    for i, heading in enumerate(is_heading):
        if heading:
            keep[i] = True
            if i + 1 < len(blocks) and not is_heading[i + 1]:
                keep[i + 1] = True

    body_budget = budget - len(_COMPRESSION_NOTICE)
    selected = []
    size = 0
    for i, block in enumerate(blocks):
        if not keep[i]:
            continue
        cost = len(block) + 2  # account for the "\n\n" joiner
        if size + cost > body_budget:
            break
        selected.append(block)
        size += cost

    result = "\n\n".join(selected).strip()
    if not result:  # pathological single huge block — hard truncate
        result = text[:body_budget].rstrip()

    return result + _COMPRESSION_NOTICE
